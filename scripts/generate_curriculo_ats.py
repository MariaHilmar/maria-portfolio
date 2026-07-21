# -*- coding: utf-8 -*-
"""Gera currículo ATS-friendly em Word (.docx)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"C:\Users\maria\OneDrive\Área de Trabalho\MariaHilmar_Curriculo_ATS.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)


def set_run_font(run, size=10.5, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(
    text,
    size=10.5,
    bold=False,
    space_before=0,
    space_after=4,
    align=None,
    color=None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.1
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_ats(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    set_run_font(run, size=11, bold=True, color=(30, 58, 95))
    return p


def add_job_header(title, company, period):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    r1 = p.add_run(title)
    set_run_font(r1, size=10.5, bold=True)
    r2 = p.add_run(f" | {company}")
    set_run_font(r2, size=10.5, bold=True, color=(30, 58, 95))
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.05
    r3 = p2.add_run(period)
    set_run_font(r3, size=9.5, color=(68, 68, 68))


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.left_indent = Inches(0.2)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=10)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10)


style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

# HEADER
add_paragraph(
    "Maria Hilmar Gomes da Silva",
    size=16,
    bold=True,
    space_after=2,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_paragraph(
    "Product Manager / Product Owner | Dados e IA | Base técnica em Python",
    size=11,
    bold=True,
    space_after=3,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    color=(30, 58, 95),
)
add_paragraph(
    "João Pessoa, PB | Disponível para atuação remota",
    size=9.5,
    space_after=1,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_paragraph(
    "Telefone: (61) 98206-2117 | E-mail: mariahilmar@gmail.com",
    size=9.5,
    space_after=1,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_paragraph(
    "LinkedIn: linkedin.com/in/mariahilmar | GitHub: github.com/MariaHilmar | "
    "Portfolio: mariahilmar.vercel.app",
    size=9.5,
    space_after=2,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)

# RESUMO
add_heading_ats("Resumo Profissional")
add_paragraph(
    "Líder de Produto com senioridade consolidada na gestão de squads ágeis e produtos digitais "
    "de alta complexidade (SaaS B2B e sistemas governamentais). Conecto negócios, arquitetura e "
    "dados: traduzo regras e legislações densas em roadmaps, requisitos rastreáveis e entregas "
    "orientadas a ROI. MBA em Ciência de Dados e estudo prático contínuo em Python, FastAPI, "
    "pipelines e fundamentos de IA, com portfólio técnico público (JurisSync, MGI KPI).",
    size=10.5,
    space_after=2,
)

# COMPETENCIAS
add_heading_ats("Competências Técnicas")
add_paragraph(
    "Gestão de Produto: Continuous Product Discovery, MVP, roadmapping, backlog orientado a "
    "valor (ROI), OKRs, KPIs, Scrum, Kanban, Lean, Management 3.0.",
    size=10,
    space_after=2,
)
add_paragraph(
    "Requisitos e Qualidade: User Stories em BDD (Gherkin), BPMN, UML, C4 Model, Análise de "
    "Pontos de Função (APF), estratégia de QA e testes (CTAL-TM), Figma, Miro.",
    size=10,
    space_after=2,
)
add_paragraph(
    "Dados e IA: produtos de dados, ETL, SQL, métricas de negócio e engenharia; estudo prático "
    "em Machine Learning (XGBoost), NLP, busca vetorial (pgvector), LLM e RAG.",
    size=10,
    space_after=2,
)
add_paragraph(
    "Stack em prática: Python, FastAPI, PostgreSQL, Supabase, Redis, Docker, Next.js, Git, CI/CD.",
    size=10,
    space_after=2,
)

# EXPERIENCIA
add_heading_ats("Experiência Profissional")

add_job_header(
    "Scrum Master / Analista de Sistemas Sênior",
    "Qintess",
    "Abril de 2025 até o presente | Remoto",
)
add_bullet(
    "Facilitação ágil (Scrum e Kanban), remoção de impedimentos e acompanhamento de "
    "Velocity e Lead Time."
)
add_bullet(
    "Refinamento de backlogs em plataformas governamentais críticas: eNatJus 4.0, "
    "Portal de Arrecadação e SAPRE (TJBA e MGI)."
)
add_bullet(
    "Tradução de regulamentações em User Stories com BDD, BPMN/UML e prototipação em Figma."
)
add_bullet(
    "Aplicação de Análise de Pontos de Função (APF) para medição funcional e conformidade "
    "com critérios de auditoria pública."
)

add_job_header(
    "Gerente de Produtos / Product Owner Sênior",
    "DFimoveis.com (Timipro)",
    "Fevereiro de 2016 a Janeiro de 2025 | Remoto",
)
add_bullet(
    "Estruturação do portfólio digital do zero, com consolidação de 6 produtos SaaS e "
    "Real Estate de alta volumetria, do Discovery ao Go-live."
)
add_bullet(
    "Liderança de squads multifuncionais sob Scrum e Kanban, com taxa de retenção de "
    "talentos técnicos acima de 90%."
)
add_bullet(
    "Implantação de OKRs com adesão corporativa e cerca de 95% de entrega histórica "
    "dos resultados-chave."
)
add_bullet(
    "Concepção de produto de inteligência analítica para predição de preços, com "
    "pipelines ETL (Azure Data Factory), SQL e dashboards no portal."
)
add_bullet(
    "Integração do ecossistema de dados e CRM via HubSpot entre Vendas, Financeiro, "
    "Customer Success e TI, com redução de cerca de 15% no churn."
)
add_bullet(
    "Iniciativas de SEO com aumento de 15% nos acessos orgânicos e 20% na conversão "
    "de leads; validação de APIs REST e testes ponta a ponta."
)

add_job_header(
    "Gerente de Produtos / Product Owner",
    "Wimoveis",
    "Janeiro de 2012 a Maio de 2015 | Brasília, DF",
)
add_bullet(
    "Modernização de sistemas fiscais e financeiros (NF-e, NFS-e, faturamento, contas "
    "a pagar/receber), com redução de cerca de 95% nos erros de conciliação."
)
add_bullet(
    "Liderança de time multifuncional e BI (SQL Server, Excel Avançado/Power Pivot), "
    "reduzindo cerca de 30% o tempo de atendimento (SLA)."
)
add_bullet(
    "Levantamento e documentação de requisitos complexos com Casos de Uso, UML e BPMN."
)

add_job_header(
    "Analista de Requisitos / Garantia de Qualidade (QA)",
    "Politec, Linkdata e Cast",
    "Janeiro de 2002 a Dezembro de 2011 | Brasília, DF",
)
add_bullet(
    "Atuação em QA e requisitos em projetos de grande porte para Banco do Brasil, "
    "Petrobras e CNJ, incluindo pioneirismo na padronização de testes no Banco do "
    "Brasil (Cast)."
)

# PROJETOS (estudo / produto proprio - NAO e emprego)
add_heading_ats("Projetos de Produto e Portfólio Técnico")
add_job_header(
    "Product Owner / Product Lead (projeto próprio de estudo)",
    "Situação Jurídica",
    "Janeiro de 2025 até o presente | Projeto pessoal | Evidências: github.com/MariaHilmar",
)
add_paragraph(
    "Produto próprio de jurimetria e dados jurídicos, usado para consolidar autonomia "
    "técnica em Python, APIs, pipelines e fundamentos de IA (não é vínculo empregatício).",
    size=9.5,
    space_after=2,
    color=(68, 68, 68),
)
add_bullet(
    "Desenho de arquitetura e implementação hands-on com FastAPI, PostgreSQL, Redis e Next.js."
)
add_bullet(
    "Exploração de dados e IA: modelos (XGBoost), NLP, busca vetorial (pgvector) e "
    "estratégias LLM/RAG aplicadas à jurimetria."
)
add_bullet(
    "Pipeline Staging/ETL versus Serving (Supabase), com redução de aproximadamente "
    "60% na latência das consultas no ambiente do projeto."
)
add_bullet(
    "Documentação técnica e critérios BDD gerados com apoio de IA, reduzindo cerca de "
    "70% o esforço de especificação no ciclo do projeto."
)
add_bullet(
    "Modelagem com isolamento multi-tenant (Row-Level Security) e premissas de "
    "Privacy by Design / LGPD."
)
add_bullet(
    "Artefatos públicos relacionados: JurisSync (API + dashboard) e MGI KPI "
    "(pipeline + BI) no GitHub e em mariahilmar.vercel.app."
)

# FORMACAO
add_heading_ats("Formação Acadêmica")
add_paragraph(
    "MBA em Ciência de Dados: Business Intelligence, Big Data e Analytics - "
    "Faculdade Anhanguera - Concluído em 2025",
    size=10,
    space_after=2,
)
add_paragraph(
    "MBA em Gestão de Projetos e Metodologias Ágeis - Faculdade Anhanguera - "
    "Concluído em 2025",
    size=10,
    space_after=2,
)
add_paragraph(
    "MBA em Teste de Software - Unieuro - Concluído em 2012",
    size=10,
    space_after=2,
)
add_paragraph(
    "Graduação em Análise de Sistemas - Instituto Nossa Senhora de Fátima - "
    "Concluído em 2011",
    size=10,
    space_after=2,
)

# CERTIFICACOES
add_heading_ats("Certificações")
add_paragraph(
    "CSPO (Certified Product Owner) - Scrum Alliance",
    size=10,
    space_after=1,
)
add_paragraph(
    "PSM I (Professional Scrum Master) - Scrum.org",
    size=10,
    space_after=1,
)
add_paragraph(
    "CTAL-TM (Certified Tester Advanced Level - Test Manager) - ISTQB",
    size=10,
    space_after=1,
)
add_paragraph("Management 3.0", size=10, space_after=1)

doc.save(OUT_PATH)
print("OK:", OUT_PATH)
print("Size:", os.path.getsize(OUT_PATH), "bytes")
